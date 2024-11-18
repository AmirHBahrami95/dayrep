import docx
import os
from utils import get_today_str,mk_file,get_settings

def __make_row(doc,r):
	is_done="[ ]"
	if r[2]:
		is_done="[x]"
	doc.add_paragraph( "{} {}".format(is_done,r[0]))

def __make_docx(rows):
	doc=docx.Document()
	doc.add_heading('DayRep ({})'.format(get_today_str().replace('-','/')),0)
	for r in rows:
		__make_row(doc,r)
	return doc

def docx_export(rows,path='today'):
	doc=__make_docx(rows)
	settings=get_settings()
	if path=='today':
		path="{}{}{}.docx".format(settings['reports_dir'],os.sep,get_today_str()).replace('-','_') # TODO make this appear in settings
	elif not path.endswith('.docx'):
		print("docx path must end with '.docx'\naborting...")
		return False
	mk_file(path)
	doc.save(path)
	return True
